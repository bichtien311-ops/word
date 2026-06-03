"""
Word Injector — Режим 2: Добавление контента в существующий документ.

Открывает существующий .docx и вставляет новый контент:
- В конец документа
- После указанного заголовка
- Перед указанным заголовком
- Замена содержимого раздела

Не ломает существующее форматирование.
Опционально — переформатирует весь документ по ГОСТ.
"""

import copy
import os
import sys
from enum import Enum
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

sys.path.insert(0, os.path.dirname(__file__))
from gost_formatter import (
    GostConfig, apply_full_gost, add_formula,
    create_gost_table, add_bibliography, apply_table_style
)
from content_processor import (
    ContentBlock, ContentBlockType, structure_content,
    clean_notebooklm_output, fix_numbers
)

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "utils"))
from logger import setup_logger

log = setup_logger("word_injector")


# =============================================================================
# МОДЕЛИ
# =============================================================================

class InsertPosition(str, Enum):
    """Где вставлять контент."""
    APPEND = "append"                  # В конец документа
    AFTER_HEADING = "after_heading"    # После указанного заголовка
    BEFORE_HEADING = "before_heading"  # Перед указанным заголовком
    REPLACE_SECTION = "replace_section"  # Заменить содержимое раздела


# =============================================================================
# АНАЛИЗ ДОКУМЕНТА
# =============================================================================

def analyze_document(doc_path: str) -> dict:
    """
    Анализирует структуру существующего документа.

    Returns:
        Словарь с информацией:
        - headings: список (level, text, paragraph_index)
        - paragraph_count: количество абзацев
        - table_count: количество таблиц
        - has_toc: есть ли оглавление
        - has_bibliography: есть ли список литературы
    """
    doc = Document(doc_path)

    headings = []
    has_toc = False
    has_bibliography = False

    for i, para in enumerate(doc.paragraphs):
        # Определяем заголовки
        if para.style.name.startswith("Heading"):
            level = int(para.style.name.replace("Heading ", "")) if " " in para.style.name else 1
            headings.append({
                "level": level,
                "text": para.text.strip(),
                "index": i
            })

            # Проверяем наличие оглавления
            text_lower = para.text.strip().lower()
            if text_lower in ("содержание", "оглавление"):
                has_toc = True
            if text_lower in ("список использованных источников",
                              "список литературы", "библиография"):
                has_bibliography = True

    result = {
        "headings": headings,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "has_toc": has_toc,
        "has_bibliography": has_bibliography
    }

    log.info(f"Анализ документа {doc_path}: "
             f"{len(headings)} заголовков, {result['paragraph_count']} абзацев, "
             f"{result['table_count']} таблиц")
    return result


def find_heading_index(doc: Document, heading_text: str,
                       case_sensitive: bool = False) -> int | None:
    """
    Ищет абзац-заголовок по тексту.

    Returns:
        Индекс абзаца или None если не найден.
    """
    search_text = heading_text if case_sensitive else heading_text.lower()

    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith("Heading"):
            para_text = para.text.strip() if case_sensitive else para.text.strip().lower()
            if search_text in para_text:
                log.info(f"Заголовок найден: '{para.text.strip()}' (индекс {i})")
                return i

    log.warning(f"Заголовок не найден: '{heading_text}'")
    return None


def find_section_range(doc: Document, heading_index: int) -> tuple[int, int]:
    """
    Определяет диапазон абзацев, принадлежащих разделу.
    Раздел = от заголовка до следующего заголовка того же или более высокого уровня.

    Returns:
        (start_index, end_index) — не включая end_index
    """
    heading_para = doc.paragraphs[heading_index]
    heading_level = _get_heading_level(heading_para)

    start = heading_index + 1  # Контент начинается после заголовка
    end = len(doc.paragraphs)   # По умолчанию — до конца документа

    for i in range(start, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style.name.startswith("Heading"):
            level = _get_heading_level(para)
            if level <= heading_level:
                end = i
                break

    log.info(f"Диапазон раздела: абзацы {start}-{end}")
    return start, end


def _get_heading_level(para) -> int:
    """Извлекает уровень заголовка из абзаца."""
    if " " in para.style.name:
        try:
            return int(para.style.name.split(" ")[-1])
        except ValueError:
            return 1
    return 1


# =============================================================================
# ВСТАВКА КОНТЕНТА
# =============================================================================

def inject_content(
    doc_path: str,
    raw_content: str,
    output_path: str,
    position: InsertPosition = InsertPosition.APPEND,
    target_heading: str = None,
    apply_gost: bool = False
) -> str:
    """
    Вставляет контент в существующий документ.

    Args:
        doc_path: Путь к исходному .docx
        raw_content: Raw-текст для вставки
        output_path: Путь для сохранения (оригинал НЕ изменяется!)
        position: Где вставлять
        target_heading: Текст заголовка (для AFTER/BEFORE/REPLACE)
        apply_gost: Применить ли ГОСТ-форматирование ко всему документу

    Returns:
        Путь к сохранённому файлу
    """
    log.info(f"=== Инъекция контента в {doc_path} ===")
    log.info(f"Позиция: {position.value}, целевой заголовок: {target_heading}")

    doc = Document(doc_path)

    # Очищаем и структурируем контент
    cleaned = clean_notebooklm_output(raw_content)
    blocks = structure_content(cleaned)

    if position == InsertPosition.APPEND:
        _inject_at_end(doc, blocks)

    elif position == InsertPosition.AFTER_HEADING:
        if not target_heading:
            raise ValueError("target_heading обязателен для AFTER_HEADING")
        idx = find_heading_index(doc, target_heading)
        if idx is None:
            raise ValueError(f"Заголовок не найден: '{target_heading}'")
        _inject_after_index(doc, blocks, idx)

    elif position == InsertPosition.BEFORE_HEADING:
        if not target_heading:
            raise ValueError("target_heading обязателен для BEFORE_HEADING")
        idx = find_heading_index(doc, target_heading)
        if idx is None:
            raise ValueError(f"Заголовок не найден: '{target_heading}'")
        _inject_before_index(doc, blocks, idx)

    elif position == InsertPosition.REPLACE_SECTION:
        if not target_heading:
            raise ValueError("target_heading обязателен для REPLACE_SECTION")
        idx = find_heading_index(doc, target_heading)
        if idx is None:
            raise ValueError(f"Заголовок не найден: '{target_heading}'")
        _replace_section(doc, blocks, idx)

    # Опциональное ГОСТ-форматирование
    if apply_gost:
        apply_full_gost(doc)
        log.info("ГОСТ-форматирование применено ко всему документу")

    # Сохранение
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    log.info(f"=== Документ сохранён: {output_path} ===")
    return output_path


def _inject_at_end(doc: Document, blocks: list[ContentBlock]) -> None:
    """Вставляет блоки в конец документа."""
    log.info(f"Вставка {len(blocks)} блоков в конец документа")
    _write_blocks(doc, blocks)


def _inject_after_index(doc: Document, blocks: list[ContentBlock],
                        heading_index: int) -> None:
    """
    Вставляет блоки после указанного заголовка.
    Находит конец текущего раздела и вставляет перед следующим заголовком.
    """
    _, section_end = find_section_range(doc, heading_index)
    log.info(f"Вставка после заголовка (индекс {heading_index}), "
             f"в позицию {section_end}")

    # Вставляем через XML-манипуляцию для точного позиционирования
    body = doc.element.body
    reference_element = doc.paragraphs[section_end - 1]._element if section_end > 0 else None

    _insert_blocks_after_element(doc, blocks, reference_element)


def _inject_before_index(doc: Document, blocks: list[ContentBlock],
                         heading_index: int) -> None:
    """Вставляет блоки перед указанным заголовком."""
    log.info(f"Вставка перед заголовком (индекс {heading_index})")

    body = doc.element.body
    reference_element = doc.paragraphs[heading_index]._element

    _insert_blocks_before_element(doc, blocks, reference_element)


def _replace_section(doc: Document, blocks: list[ContentBlock],
                     heading_index: int) -> None:
    """
    Заменяет содержимое раздела (но сохраняет сам заголовок).
    """
    start, end = find_section_range(doc, heading_index)
    log.info(f"Замена раздела: удаление абзацев {start}-{end}")

    # Удаляем старые абзацы раздела (в обратном порядке)
    body = doc.element.body
    elements_to_remove = []
    for i in range(start, min(end, len(doc.paragraphs))):
        elements_to_remove.append(doc.paragraphs[i]._element)

    for elem in elements_to_remove:
        body.remove(elem)

    # Вставляем новый контент после заголовка
    heading_element = doc.paragraphs[heading_index]._element
    _insert_blocks_after_element(doc, blocks, heading_element)

    log.info(f"Раздел заменён: {len(elements_to_remove)} абзацев удалено, "
             f"{len(blocks)} блоков вставлено")


# =============================================================================
# НИЗКОУРОВНЕВАЯ ВСТАВКА
# =============================================================================

def _write_blocks(doc: Document, blocks: list[ContentBlock]) -> None:
    """Записывает блоки контента в конец документа (простой режим)."""
    table_counter = _count_existing_tables(doc)
    formula_counter = _count_existing_formulas(doc)

    for block in blocks:
        if block.block_type == ContentBlockType.HEADING:
            level = max(1, min(block.level, 3))
            text = block.content
            if block.number:
                text = f"{block.number} {text}"
            doc.add_heading(text, level=level)

        elif block.block_type == ContentBlockType.PARAGRAPH:
            doc.add_paragraph(block.content, style="Normal")

        elif block.block_type == ContentBlockType.FORMULA:
            formula_counter += 1
            plain_text = block.metadata.get("plain_text", block.content)
            variables = block.metadata.get("variables", None)
            add_formula(doc, plain_text, number=str(formula_counter), variables=variables)

        elif block.block_type == ContentBlockType.TABLE:
            table_counter += 1
            headers = block.metadata.get("headers", [])
            rows = block.metadata.get("rows", [])
            caption = block.metadata.get("caption", "")
            if headers and rows:
                create_gost_table(doc, headers, rows,
                                  caption=caption, number=str(table_counter))

        elif block.block_type == ContentBlockType.LIST_ITEM:
            para = doc.add_paragraph(block.content, style="List Bullet")
            for run in para.runs:
                run.font.name = GostConfig.FONT_NAME
                run.font.size = GostConfig.FONT_SIZE


def _insert_blocks_after_element(doc: Document, blocks: list[ContentBlock],
                                 reference_element) -> None:
    """Вставляет блоки после указанного XML-элемента."""
    # Создаём временный документ для генерации элементов
    temp_doc = Document()
    _write_blocks(temp_doc, blocks)

    body = doc.element.body

    # Вставляем элементы из temp_doc после reference_element
    insert_after = reference_element
    for child in list(temp_doc.element.body):
        new_elem = copy.deepcopy(child)
        if insert_after is not None:
            insert_after.addnext(new_elem)
        else:
            body.append(new_elem)
        insert_after = new_elem


def _insert_blocks_before_element(doc: Document, blocks: list[ContentBlock],
                                  reference_element) -> None:
    """Вставляет блоки перед указанным XML-элементом."""
    temp_doc = Document()
    _write_blocks(temp_doc, blocks)

    for child in reversed(list(temp_doc.element.body)):
        new_elem = copy.deepcopy(child)
        reference_element.addprevious(new_elem)


# =============================================================================
# УТИЛИТЫ
# =============================================================================

def _count_existing_tables(doc: Document) -> int:
    """Подсчитывает существующие таблицы для продолжения нумерации."""
    return len(doc.tables)


def _count_existing_formulas(doc: Document) -> int:
    """Подсчитывает существующие формулы (по паттерну нумерации) для продолжения."""
    import re
    count = 0
    for para in doc.paragraphs:
        if re.search(r"\(\d+\)", para.text):
            # Эвристика: абзац с номером формулы в скобках
            count += 1
    return count


def reformat_existing_document(doc_path: str, output_path: str) -> str:
    """
    Переформатирует существующий документ по ГОСТ.
    Контент не меняется, только стили.

    Args:
        doc_path: Путь к исходному .docx
        output_path: Путь для сохранения

    Returns:
        Путь к сохранённому файлу
    """
    log.info(f"Переформатирование: {doc_path}")
    doc = Document(doc_path)
    apply_full_gost(doc)

    # Применяем TNR ко всем абзацам
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name is None or run.font.name != GostConfig.FONT_NAME:
                run.font.name = GostConfig.FONT_NAME
            if run.font.size is None:
                run.font.size = GostConfig.FONT_SIZE

    # Применяем стиль к таблицам
    for table in doc.tables:
        apply_table_style(table)

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    log.info(f"Документ переформатирован: {output_path}")
    return output_path


# =============================================================================
# CLI
# =============================================================================

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Word Injector — вставка контента в существующий документ",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры:
  python word_injector.py --target doc.docx --input new_content.txt --output result.docx --position append
  python word_injector.py --target doc.docx --input content.txt --output result.docx --position after_heading --heading "Введение"
  python word_injector.py --reformat doc.docx --output formatted.docx
  python word_injector.py --analyze doc.docx
        """
    )
    parser.add_argument("--target", help="Путь к существующему .docx")
    parser.add_argument("--input", help="Путь к файлу с контентом для вставки")
    parser.add_argument("--output", required=True, help="Путь для сохранения")
    parser.add_argument("--position", choices=[p.value for p in InsertPosition],
                        default="append", help="Где вставлять")
    parser.add_argument("--heading", help="Целевой заголовок (для after/before/replace)")
    parser.add_argument("--apply-gost", action="store_true", help="Применить ГОСТ ко всему документу")
    parser.add_argument("--reformat", metavar="DOCX", help="Только переформатировать по ГОСТ")
    parser.add_argument("--analyze", metavar="DOCX", help="Только проанализировать структуру")

    args = parser.parse_args()

    if args.analyze:
        info = analyze_document(args.analyze)
        print(f"Абзацев: {info['paragraph_count']}")
        print(f"Таблиц: {info['table_count']}")
        print(f"Оглавление: {'да' if info['has_toc'] else 'нет'}")
        print(f"Список литературы: {'да' if info['has_bibliography'] else 'нет'}")
        print(f"\nЗаголовки ({len(info['headings'])}):")
        for h in info["headings"]:
            indent = "  " * h["level"]
            print(f"  {indent}[H{h['level']}] {h['text']}")

    elif args.reformat:
        reformat_existing_document(args.reformat, args.output)
        print(f"✓ Документ переформатирован: {args.output}")

    elif args.target and args.input:
        with open(args.input, "r", encoding="utf-8") as f:
            raw_content = f.read()
        inject_content(
            doc_path=args.target,
            raw_content=raw_content,
            output_path=args.output,
            position=InsertPosition(args.position),
            target_heading=args.heading,
            apply_gost=args.apply_gost
        )
        print(f"[OK] Успешно! Файл сохранён: {args.output}")
    else:
        parser.print_help()
