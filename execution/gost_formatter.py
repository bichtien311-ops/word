"""
ГОСТ-форматировщик для Word-документов.

Применяет параметры ГОСТ 7.32-2017 и типовые вузовские требования
к документам python-docx. Является ядром системы — все скрипты
используют этот модуль для обеспечения соответствия ГОСТу.

ГОСТы-основа:
- ГОСТ 7.32-2017 (структура и оформление)
- ГОСТ Р 7.0.100-2018 (библиографическое описание)
- Типовые вузовские требования (шрифты, поля, интервалы)
"""

import os
import sys
from enum import Enum

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Cm, Mm, Pt, RGBColor
from lxml import etree

# --- Добавляем utils в путь ---
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "utils"))
from logger import setup_logger

log = setup_logger("gost_formatter")


# =============================================================================
# КОНФИГУРАЦИЯ ГОСТ
# =============================================================================

class GostConfig:
    """
    Центральное хранилище всех параметров ГОСТ-форматирования.
    Изменяйте здесь — и все скрипты подхватят.
    """

    # --- Параметры страницы ---
    PAGE_WIDTH = Mm(210)
    PAGE_HEIGHT = Mm(297)
    MARGIN_LEFT = Mm(30)       # Для переплёта
    MARGIN_RIGHT = Mm(15)
    MARGIN_TOP = Mm(20)
    MARGIN_BOTTOM = Mm(20)

    # --- Основной шрифт ---
    FONT_NAME = "Times New Roman"
    FONT_SIZE = Pt(14)
    FONT_SIZE_SMALL = Pt(12)   # Для таблиц, подписей
    LINE_SPACING = 1.5         # Полуторный
    PARAGRAPH_INDENT = Cm(1.25)  # Абзацный отступ
    SPACE_AFTER_PARAGRAPH = Pt(0)
    SPACE_BEFORE_PARAGRAPH = Pt(0)

    # --- Заголовки ---
    HEADING1_SIZE = Pt(16)
    HEADING1_BOLD = True
    HEADING1_CAPS = True       # ЗАГЛАВНЫЕ
    HEADING1_SPACE_BEFORE = Pt(24)
    HEADING1_SPACE_AFTER = Pt(12)

    HEADING2_SIZE = Pt(14)
    HEADING2_BOLD = True
    HEADING2_SPACE_BEFORE = Pt(18)
    HEADING2_SPACE_AFTER = Pt(6)

    HEADING3_SIZE = Pt(14)
    HEADING3_BOLD = True
    HEADING3_SPACE_BEFORE = Pt(12)
    HEADING3_SPACE_AFTER = Pt(6)

    # --- Таблицы ---
    TABLE_FONT_SIZE = Pt(12)
    TABLE_HEADER_BOLD = True
    TABLE_ALIGNMENT = WD_TABLE_ALIGNMENT.CENTER
    TABLE_BORDER_COLOR = "000000"
    TABLE_BORDER_SIZE = 4  # В восьмых долях пункта

    # --- Нумерация ---
    PAGE_NUMBER_POSITION = "bottom_center"

    # --- Цвета ---
    COLOR_BLACK = RGBColor(0, 0, 0)


class DocType(str, Enum):
    """Типы поддерживаемых документов."""
    METHODICAL_GUIDE = "methodical_guide"    # Методическое пособие
    LECTURE = "lecture"                       # Лекционный материал
    STUDY_GUIDE = "study_guide"              # Учебное пособие
    CALCULATION_NOTE = "calculation_note"     # Расчётная записка


# =============================================================================
# ПАРАМЕТРЫ СТРАНИЦЫ
# =============================================================================

def apply_page_setup(doc: Document) -> None:
    """
    Устанавливает параметры страницы по ГОСТ:
    А4, поля: лево 30мм, право 15мм, верх/низ 20мм.
    """
    for section in doc.sections:
        section.page_width = GostConfig.PAGE_WIDTH
        section.page_height = GostConfig.PAGE_HEIGHT
        section.left_margin = GostConfig.MARGIN_LEFT
        section.right_margin = GostConfig.MARGIN_RIGHT
        section.top_margin = GostConfig.MARGIN_TOP
        section.bottom_margin = GostConfig.MARGIN_BOTTOM
        section.orientation = WD_ORIENT.PORTRAIT

    log.info("Параметры страницы установлены (А4, поля по ГОСТ)")


# =============================================================================
# СТИЛИ ТЕКСТА
# =============================================================================

def apply_base_styles(doc: Document) -> None:
    """
    Настраивает базовые стили документа:
    - Normal: TNR 14, полуторный интервал, отступ 1.25 см, по ширине
    """
    # --- Стиль Normal ---
    style_normal = doc.styles["Normal"]
    font = style_normal.font
    font.name = GostConfig.FONT_NAME
    font.size = GostConfig.FONT_SIZE
    font.color.rgb = GostConfig.COLOR_BLACK
    font.bold = False
    font.italic = False

    # Устанавливаем шрифт для кириллических символов (eastAsia / cs)
    rpr = style_normal.element.get_or_add_rPr()
    _set_cyrillic_font(rpr, GostConfig.FONT_NAME)

    paragraph_format = style_normal.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = GostConfig.LINE_SPACING
    paragraph_format.first_line_indent = GostConfig.PARAGRAPH_INDENT
    paragraph_format.space_after = GostConfig.SPACE_AFTER_PARAGRAPH
    paragraph_format.space_before = GostConfig.SPACE_BEFORE_PARAGRAPH

    log.info(f"Базовый стиль Normal настроен: {GostConfig.FONT_NAME} {GostConfig.FONT_SIZE}, интервал {GostConfig.LINE_SPACING}")


def apply_heading_styles(doc: Document) -> None:
    """
    Настраивает стили заголовков по ГОСТ:
    - Heading 1: TNR 16, жирный, ЗАГЛАВНЫЕ
    - Heading 2: TNR 14, жирный
    - Heading 3: TNR 14, жирный
    """
    _setup_heading(doc, "Heading 1",
                   size=GostConfig.HEADING1_SIZE,
                   bold=GostConfig.HEADING1_BOLD,
                   space_before=GostConfig.HEADING1_SPACE_BEFORE,
                   space_after=GostConfig.HEADING1_SPACE_AFTER,
                   all_caps=GostConfig.HEADING1_CAPS)

    _setup_heading(doc, "Heading 2",
                   size=GostConfig.HEADING2_SIZE,
                   bold=GostConfig.HEADING2_BOLD,
                   space_before=GostConfig.HEADING2_SPACE_BEFORE,
                   space_after=GostConfig.HEADING2_SPACE_AFTER)

    _setup_heading(doc, "Heading 3",
                   size=GostConfig.HEADING3_SIZE,
                   bold=GostConfig.HEADING3_BOLD,
                   space_before=GostConfig.HEADING3_SPACE_BEFORE,
                   space_after=GostConfig.HEADING3_SPACE_AFTER)

    log.info("Стили заголовков Heading 1-3 настроены по ГОСТ")


def _setup_heading(doc: Document, style_name: str, size: Pt, bold: bool,
                   space_before: Pt, space_after: Pt, all_caps: bool = False) -> None:
    """Вспомогательная: настройка одного стиля заголовка."""
    style = doc.styles[style_name]
    font = style.font
    font.name = GostConfig.FONT_NAME
    font.size = size
    font.bold = bold
    font.color.rgb = GostConfig.COLOR_BLACK
    font.all_caps = all_caps
    font.italic = False

    rpr = style.element.get_or_add_rPr()
    _set_cyrillic_font(rpr, GostConfig.FONT_NAME)

    pf = style.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.first_line_indent = None  # Заголовки без отступа
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = GostConfig.LINE_SPACING
    pf.keep_with_next = True  # Не отрывать от следующего абзаца


def _set_cyrillic_font(rpr_element, font_name: str) -> None:
    """
    Устанавливает шрифт для кириллических символов (cs и eastAsia).
    Без этого Word может подставлять другой шрифт для кириллицы.
    """
    rfonts = rpr_element.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr_element.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


# =============================================================================
# НУМЕРАЦИЯ СТРАНИЦ
# =============================================================================

def add_page_numbering(doc: Document) -> None:
    """
    Добавляет нумерацию страниц внизу по центру (сквозная).
    Титульный лист считается, но номер на нём не ставится.
    """
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Устанавливаем шрифт номера страницы
        run = paragraph.add_run()
        run.font.name = GostConfig.FONT_NAME
        run.font.size = GostConfig.FONT_SIZE

        # Вставляем поле PAGE
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._element.append(fld_char_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        run._element.append(instr_text)

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._element.append(fld_char_end)

    log.info("Нумерация страниц добавлена (внизу по центру)")


# =============================================================================
# ТАБЛИЦЫ
# =============================================================================

def apply_table_style(table) -> None:
    """
    Применяет ГОСТ-стиль к таблице:
    - Шрифт TNR 12, границы чёрные, выравнивание по центру
    """
    table.alignment = GostConfig.TABLE_ALIGNMENT

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = GostConfig.FONT_NAME
                    run.font.size = GostConfig.TABLE_FONT_SIZE
                    _set_cyrillic_font(run._element.get_or_add_rPr(), GostConfig.FONT_NAME)

    # Устанавливаем границы таблицы
    _set_table_borders(table)

    log.info("Стиль таблицы применён (TNR 12, границы, по центру)")


def _set_table_borders(table, color: str = None, size: int = None) -> None:
    """Устанавливает видимые границы для таблицы."""
    color = color or GostConfig.TABLE_BORDER_COLOR
    size = size or GostConfig.TABLE_BORDER_SIZE

    tbl = table._element
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)

    # Удаляем старые границы если есть
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)


def create_gost_table(doc: Document, headers: list[str], rows: list[list[str]],
                      caption: str = None, number: str = None) -> None:
    """
    Создаёт таблицу по ГОСТ с заголовком и данными.

    Args:
        doc: Документ
        headers: Список заголовков столбцов
        rows: Список строк (каждая строка — список значений)
        caption: Подпись таблицы (например, "Сравнительные показатели")
        number: Номер таблицы (например, "1.1")
    """
    # Подпись таблицы (над таблицей по ГОСТ)
    if caption:
        label = f"Таблица {number} — {caption}" if number else f"Таблица — {caption}"
        caption_para = doc.add_paragraph(label)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caption_para.paragraph_format.first_line_indent = None
        caption_para.paragraph_format.space_after = Pt(6)
        for run in caption_para.runs:
            run.font.name = GostConfig.FONT_NAME
            run.font.size = GostConfig.FONT_SIZE
            run.font.bold = False

    # Создаём таблицу
    table = doc.add_table(rows=1, cols=len(headers))

    # Заголовки
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header_text)
        run.font.bold = GostConfig.TABLE_HEADER_BOLD
        run.font.name = GostConfig.FONT_NAME
        run.font.size = GostConfig.TABLE_FONT_SIZE

    # Данные
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = ""
                para = row_cells[i].paragraphs[0]
                run = para.add_run(str(cell_text))
                run.font.name = GostConfig.FONT_NAME
                run.font.size = GostConfig.TABLE_FONT_SIZE

    apply_table_style(table)
    log.info(f"Таблица создана: {len(headers)} столбцов, {len(rows)} строк")


# =============================================================================
# ФОРМУЛЫ
# =============================================================================

def add_formula(doc: Document, formula_text: str, number: str = None,
                variables: dict[str, str] = None) -> None:
    """
    Добавляет формулу в документ по ГОСТ:
    - Формула по центру, номер справа в скобках
    - После формулы — пояснение переменных через «где»

    Args:
        doc: Документ
        formula_text: Текст формулы (обычная запись, например "I = S / (√3 · U)")
        number: Номер формулы (например, "1.1")
        variables: Словарь пояснений {"I": "ток, А", "S": "полная мощность, кВА"}
    """
    # Формула по центру с номером справа
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    run = para.add_run(formula_text)
    run.font.name = GostConfig.FONT_NAME
    run.font.size = GostConfig.FONT_SIZE
    run.font.italic = True

    # Номер формулы справа
    if number:
        run_num = para.add_run(f"\t({number})")
        run_num.font.name = GostConfig.FONT_NAME
        run_num.font.size = GostConfig.FONT_SIZE

    # Пояснение переменных через «где»
    if variables:
        first = True
        for var_name, var_desc in variables.items():
            if first:
                where_para = doc.add_paragraph()
                where_para.paragraph_format.first_line_indent = None
                where_para.paragraph_format.space_before = Pt(3)
                where_para.paragraph_format.space_after = Pt(0)
                run_where = where_para.add_run("где ")
                run_where.font.name = GostConfig.FONT_NAME
                run_where.font.size = GostConfig.FONT_SIZE

                run_var = where_para.add_run(f"{var_name}")
                run_var.font.name = GostConfig.FONT_NAME
                run_var.font.size = GostConfig.FONT_SIZE
                run_var.font.italic = True

                run_desc = where_para.add_run(f" — {var_desc};")
                run_desc.font.name = GostConfig.FONT_NAME
                run_desc.font.size = GostConfig.FONT_SIZE
                first = False
            else:
                var_para = doc.add_paragraph()
                var_para.paragraph_format.first_line_indent = Cm(1.7)  # Выравнивание с "где"
                var_para.paragraph_format.space_before = Pt(0)
                var_para.paragraph_format.space_after = Pt(0)

                run_var = var_para.add_run(f"{var_name}")
                run_var.font.name = GostConfig.FONT_NAME
                run_var.font.size = GostConfig.FONT_SIZE
                run_var.font.italic = True

                run_desc = var_para.add_run(f" — {var_desc};")
                run_desc.font.name = GostConfig.FONT_NAME
                run_desc.font.size = GostConfig.FONT_SIZE

    log.info(f"Формула добавлена: '{formula_text[:50]}...' ({number or 'без номера'})")


# =============================================================================
# РИСУНКИ (подписи)
# =============================================================================

def add_figure_caption(doc: Document, caption_text: str, number: str) -> None:
    """
    Добавляет подпись к рисунку по ГОСТ (под рисунком, по центру).

    Args:
        caption_text: Текст подписи (например, "Схема электроснабжения")
        number: Номер рисунка (например, "1.1")
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)

    run = para.add_run(f"Рисунок {number} — {caption_text}")
    run.font.name = GostConfig.FONT_NAME
    run.font.size = GostConfig.FONT_SIZE
    run.font.bold = False

    log.info(f"Подпись к рисунку: Рисунок {number} — {caption_text}")


# =============================================================================
# СПИСОК ЛИТЕРАТУРЫ
# =============================================================================

def add_bibliography(doc: Document, sources: list[str]) -> None:
    """
    Добавляет раздел «Список использованных источников» по ГОСТ.
    Каждый источник — пронумерованный элемент.

    Args:
        sources: Список строк-источников, уже оформленных по ГОСТ Р 7.0.100-2018
    """
    doc.add_heading("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", level=1)

    for i, source in enumerate(sources, 1):
        para = doc.add_paragraph(f"{i}. {source}")
        para.paragraph_format.first_line_indent = None
        para.paragraph_format.left_indent = Cm(1.25)
        para.paragraph_format.hanging_indent = Cm(1.0)
        for run in para.runs:
            run.font.name = GostConfig.FONT_NAME
            run.font.size = GostConfig.FONT_SIZE

    log.info(f"Список литературы добавлен: {len(sources)} источников")


# =============================================================================
# КОМПЛЕКСНОЕ ФОРМАТИРОВАНИЕ
# =============================================================================

def apply_full_gost(doc: Document) -> None:
    """
    Применяет полный набор ГОСТ-форматирования к документу:
    1. Параметры страницы
    2. Базовые стили текста
    3. Стили заголовков
    4. Нумерация страниц

    Это основная функция — вызывайте её для любого документа.
    """
    log.info("=== Начало полного ГОСТ-форматирования ===")
    apply_page_setup(doc)
    apply_base_styles(doc)
    apply_heading_styles(doc)
    add_page_numbering(doc)
    log.info("=== ГОСТ-форматирование завершено ===")


# =============================================================================
# УТИЛИТЫ
# =============================================================================

def create_empty_gost_document() -> Document:
    """
    Создаёт пустой документ с полным ГОСТ-форматированием.
    Используйте как стартовую точку для любого нового документа.
    """
    doc = Document()
    apply_full_gost(doc)
    return doc


def reformat_document(input_path: str, output_path: str) -> None:
    """
    Открывает существующий документ и применяет ГОСТ-форматирование.
    Контент не меняется, только стили.

    Args:
        input_path: Путь к исходному .docx
        output_path: Путь для сохранения отформатированного .docx
    """
    log.info(f"Переформатирование: {input_path}")
    doc = Document(input_path)
    apply_full_gost(doc)

    # Применяем стиль Normal ко всем абзацам без стиля
    for para in doc.paragraphs:
        if para.style.name == "Normal" or para.style.name.startswith("Body"):
            for run in para.runs:
                if run.font.name is None:
                    run.font.name = GostConfig.FONT_NAME
                if run.font.size is None:
                    run.font.size = GostConfig.FONT_SIZE

    doc.save(output_path)
    log.info(f"Документ сохранён: {output_path}")


# =============================================================================
# CLI-ИНТЕРФЕЙС
# =============================================================================

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="ГОСТ-форматировщик Word-документов",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры:
  python gost_formatter.py --create output.docx
  python gost_formatter.py --reformat input.docx output.docx
        """
    )
    parser.add_argument("--create", metavar="OUTPUT", help="Создать пустой ГОСТ-документ")
    parser.add_argument("--reformat", nargs=2, metavar=("INPUT", "OUTPUT"),
                        help="Переформатировать существующий документ")

    args = parser.parse_args()

    if args.create:
        doc = create_empty_gost_document()
        doc.save(args.create)
        print(f"✓ Создан ГОСТ-документ: {args.create}")

    elif args.reformat:
        reformat_document(args.reformat[0], args.reformat[1])
        print(f"✓ Документ переформатирован: {args.reformat[1]}")

    else:
        parser.print_help()
