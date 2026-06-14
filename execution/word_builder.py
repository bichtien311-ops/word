"""
Word Builder — Режим 1: Создание документа с нуля.

Принимает структурированный контент (из content_processor) и создаёт
полноценный Word-документ по ГОСТ, включая:
- Титульный лист
- Оглавление
- Основная часть (разделы, таблицы, формулы)
- Список литературы
- Приложения

Использует gost_formatter.py для всего форматирования.
"""

import os
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

# --- Импорт внутренних модулей ---
sys.path.insert(0, os.path.dirname(__file__))
from gost_formatter import (
    GostConfig, DocType, apply_full_gost, add_formula,
    create_gost_table, add_bibliography, add_figure_caption
)
from content_processor import (
    ContentBlock, ContentBlockType, structure_content,
    clean_notebooklm_output, fix_numbers
)

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "utils"))
from logger import setup_logger

log = setup_logger("word_builder")


# =============================================================================
# МОДЕЛИ ДАННЫХ
# =============================================================================

@dataclass
class DocumentMetadata:
    """Метаданные документа для титульного листа."""
    title: str = ""                    # Название работы
    doc_type_label: str = ""           # "МЕТОДИЧЕСКОЕ ПОСОБИЕ", "ЛЕКЦИОННЫЙ МАТЕРИАЛ" и т.д.
    university: str = ""               # Название вуза
    faculty: str = ""                  # Факультет
    department: str = ""               # Кафедра
    discipline: str = ""               # Дисциплина
    author: str = ""                   # Автор (ФИО)
    author_title: str = ""             # Должность автора
    city: str = ""                     # Город
    year: str = ""                     # Год
    reviewer: str = ""                 # Рецензент (опционально)
    annotation: str = ""               # Аннотация (для учебных пособий)


@dataclass
class Section:
    """Раздел документа."""
    title: str
    level: int = 1
    content_blocks: list[ContentBlock] = field(default_factory=list)
    subsections: list = field(default_factory=list)  # Вложенные Section


# =============================================================================
# ТИТУЛЬНЫЙ ЛИСТ
# =============================================================================

def build_title_page(doc: Document, meta: DocumentMetadata) -> None:
    """
    Создаёт титульный лист по типовому образцу ВУЗа.
    Все элементы — Times New Roman, по центру.
    """
    log.info("Создание титульного листа...")

    # --- Шапка: ВУЗ ---
    if meta.university:
        _add_title_line(doc, "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ", Pt(12), bold=False)
        _add_title_line(doc, "РОССИЙСКОЙ ФЕДЕРАЦИИ", Pt(12), bold=False)
        _add_title_line(doc, "", Pt(6))  # Пустая строка
        _add_title_line(doc, meta.university.upper(), Pt(14), bold=True)

    if meta.faculty:
        _add_title_line(doc, "", Pt(6))
        _add_title_line(doc, meta.faculty, Pt(14), bold=False)

    if meta.department:
        _add_title_line(doc, meta.department, Pt(14), bold=False)

    # --- Пустое пространство перед названием ---
    for _ in range(3):
        _add_title_line(doc, "", Pt(14))

    # --- Тип документа ---
    if meta.doc_type_label:
        _add_title_line(doc, meta.doc_type_label.upper(), Pt(16), bold=True)

    # --- Название работы ---
    if meta.title:
        _add_title_line(doc, "", Pt(6))
        _add_title_line(doc, meta.title, Pt(18), bold=True)

    # --- Дисциплина ---
    if meta.discipline:
        _add_title_line(doc, "", Pt(6))
        _add_title_line(doc, f"по дисциплине «{meta.discipline}»", Pt(14), bold=False)

    # --- Пустое пространство ---
    for _ in range(4):
        _add_title_line(doc, "", Pt(14))

    # --- Автор (справа) ---
    if meta.author:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.first_line_indent = None
        para.paragraph_format.space_after = Pt(0)

        if meta.author_title:
            run = para.add_run(f"{meta.author_title}\n")
            run.font.name = GostConfig.FONT_NAME
            run.font.size = Pt(14)

        run = para.add_run(meta.author)
        run.font.name = GostConfig.FONT_NAME
        run.font.size = Pt(14)

    # --- Пустое пространство ---
    for _ in range(3):
        _add_title_line(doc, "", Pt(14))

    # --- Город и год ---
    if meta.city and meta.year:
        _add_title_line(doc, f"{meta.city} — {meta.year}", Pt(14), bold=False)
    elif meta.city:
        _add_title_line(doc, meta.city, Pt(14), bold=False)

    # --- Разрыв страницы после титульного ---
    doc.add_page_break()
    log.info("Титульный лист создан")


def _add_title_line(doc: Document, text: str, size: Pt, bold: bool = False) -> None:
    """Вспомогательная: добавляет строку на титульный лист."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0

    run = para.add_run(text)
    run.font.name = GostConfig.FONT_NAME
    run.font.size = size
    run.font.bold = bold


# =============================================================================
# ОГЛАВЛЕНИЕ
# =============================================================================

def build_table_of_contents(doc: Document) -> None:
    """
    Вставляет поле автособираемого оглавления.
    При открытии в Word нужно нажать Ctrl+A → F9 для обновления.
    """
    log.info("Вставка оглавления...")

    # Заголовок
    heading_para = doc.add_heading("СОДЕРЖАНИЕ", level=1)

    # Поле оглавления (TOC)
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = None

    run = para.add_run()

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._element.append(fld_char_separate)

    # Плейсхолдер текст
    run2 = para.add_run("Обновите оглавление: выделите всё (Ctrl+A) → обновите поля (F9)")
    run2.font.name = GostConfig.FONT_NAME
    run2.font.size = GostConfig.FONT_SIZE
    run2.font.color.rgb = GostConfig.COLOR_BLACK

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run2._element.append(fld_char_end)

    doc.add_page_break()
    log.info("Оглавление вставлено")


# =============================================================================
# ОСНОВНАЯ ЧАСТЬ
# =============================================================================

def build_sections(doc: Document, blocks: list[ContentBlock]) -> None:
    """
    Строит основную часть документа из списка блоков контента.
    Каждый блок вставляется с соответствующим ГОСТ-форматированием.
    """
    log.info(f"Построение основной части: {len(blocks)} блоков")

    table_counter = 0
    formula_counter = 0

    for block in blocks:
        if block.block_type == ContentBlockType.HEADING:
            level = max(1, min(block.level, 3))
            heading_text = block.content
            if block.number:
                heading_text = f"{block.number} {heading_text}"
            doc.add_heading(heading_text, level=level)

        elif block.block_type == ContentBlockType.PARAGRAPH:
            para = doc.add_paragraph(block.content, style="Normal")

        elif block.block_type == ContentBlockType.FORMULA:
            formula_counter += 1
            plain_text = block.metadata.get("plain_text", block.content)
            variables = block.metadata.get("variables", None)
            add_formula(doc, plain_text, number=str(formula_counter),
                        variables=variables, latex=block.content)

        elif block.block_type == ContentBlockType.TABLE:
            table_counter += 1
            headers = block.metadata.get("headers", [])
            rows = block.metadata.get("rows", [])
            caption = block.metadata.get("caption", "")
            if headers and rows:
                create_gost_table(doc, headers, rows,
                                  caption=caption, number=str(table_counter))

        elif block.block_type == ContentBlockType.LIST_ITEM:
            para = doc.add_paragraph(block.content, style="List Bullet")
            for run in para.runs:
                run.font.name = GostConfig.FONT_NAME
                run.font.size = GostConfig.FONT_SIZE

    log.info(f"Основная часть построена: {table_counter} таблиц, {formula_counter} формул")


# =============================================================================
# ВВЕДЕНИЕ / ЗАКЛЮЧЕНИЕ
# =============================================================================

def build_introduction(doc: Document, text: str) -> None:
    """Добавляет раздел «Введение»."""
    doc.add_heading("ВВЕДЕНИЕ", level=1)
    para = doc.add_paragraph(text, style="Normal")
    log.info("Введение добавлено")


def build_conclusion(doc: Document, text: str) -> None:
    """Добавляет раздел «Заключение»."""
    doc.add_heading("ЗАКЛЮЧЕНИЕ", level=1)
    para = doc.add_paragraph(text, style="Normal")
    log.info("Заключение добавлено")


# =============================================================================
# ГЛАВНАЯ ФУНКЦИЯ: СОЗДАНИЕ ДОКУМЕНТА
# =============================================================================

def create_document(
    doc_type: DocType,
    raw_content: str,
    metadata: DocumentMetadata,
    output_path: str,
    sources: list[str] = None,
    introduction: str = None,
    conclusion: str = None
) -> str:
    """
    Создаёт полный документ по ГОСТ из raw-контента.

    Пайплайн:
    1. Очистка raw-контента (NotebookLM)
    2. Структурирование в блоки
    3. Замена . → , в числах
    4. Создание документа:
       - Титульный лист
       - Оглавление
       - Введение
       - Основная часть
       - Заключение
       - Список литературы
    5. Применение ГОСТ-форматирования
    6. Сохранение

    Args:
        doc_type: Тип документа (методичка, лекция, пособие, записка)
        raw_content: Raw-текст (из NotebookLM, Markdown, или plain text)
        metadata: Метаданные (титульный лист)
        output_path: Путь для сохранения .docx
        sources: Список литературы
        introduction: Текст введения
        conclusion: Текст заключения

    Returns:
        Путь к созданному файлу
    """
    log.info(f"=== Создание документа: {doc_type.value} ===")

    # Автоподстановка типа документа
    type_labels = {
        DocType.METHODICAL_GUIDE: "МЕТОДИЧЕСКОЕ ПОСОБИЕ",
        DocType.LECTURE: "ЛЕКЦИОННЫЙ МАТЕРИАЛ",
        DocType.STUDY_GUIDE: "УЧЕБНОЕ ПОСОБИЕ",
        DocType.CALCULATION_NOTE: "РАСЧЁТНАЯ ЗАПИСКА",
    }
    if not metadata.doc_type_label:
        metadata.doc_type_label = type_labels.get(doc_type, "")

    # 1. Очистка контента
    cleaned = clean_notebooklm_output(raw_content)

    # 2. Структурирование
    blocks = structure_content(cleaned)

    # 3. Создание документа
    doc = Document()

    # 4. Титульный лист
    build_title_page(doc, metadata)

    # 5. Оглавление
    build_table_of_contents(doc)

    # 6. Введение
    if introduction:
        build_introduction(doc, fix_numbers(introduction))

    # 7. Основная часть
    build_sections(doc, blocks)

    # 8. Заключение
    if conclusion:
        build_conclusion(doc, fix_numbers(conclusion))

    # 9. Список литературы
    if sources:
        add_bibliography(doc, sources)

    # 10. Полное ГОСТ-форматирование
    apply_full_gost(doc)

    # 11. Сохранение
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    log.info(f"=== Документ сохранён: {output_path} ===")

    return output_path


# =============================================================================
# CLI
# =============================================================================

if __name__ == "__main__":
    import argparse
    from datetime import datetime

    parser = argparse.ArgumentParser(
        description="Word Builder — создание документов по ГОСТ",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры:
  python word_builder.py --type methodical_guide --input data/task.md --output .tmp/output/document.docx
  python word_builder.py --type lecture --input raw_text.txt --output lecture.docx --title "Электроснабжение"
        """
    )
    parser.add_argument("--type", choices=[t.value for t in DocType], required=True,
                        help="Тип документа")
    parser.add_argument("--input", required=True, help="Путь к файлу с raw-контентом")
    parser.add_argument("--output", required=True, help="Путь для сохранения .docx")
    parser.add_argument("--title", default="", help="Название работы")
    parser.add_argument("--university", default="", help="Название вуза")
    parser.add_argument("--author", default="", help="Автор (ФИО)")
    parser.add_argument("--discipline", default="", help="Дисциплина")
    parser.add_argument("--city", default="", help="Город")

    args = parser.parse_args()

    # Читаем контент
    with open(args.input, "r", encoding="utf-8") as f:
        raw_content = f.read()

    # Метаданные
    meta = DocumentMetadata(
        title=args.title,
        university=args.university,
        author=args.author,
        discipline=args.discipline,
        city=args.city,
        year=str(datetime.now().year)
    )

    # Создаём документ
    result = create_document(
        doc_type=DocType(args.type),
        raw_content=raw_content,
        metadata=meta,
        output_path=args.output
    )
    print(f"[OK] Документ успешно создан: {result}")
