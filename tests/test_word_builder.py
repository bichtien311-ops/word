"""
Интеграционные тесты для word_builder.py

Проверяет полный пайплайн: raw-текст → структурирование → .docx
"""

import os
import sys
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "execution"))

import pytest
from docx import Document

from word_builder import (
    DocumentMetadata,
    DocType,
    create_document,
    build_title_page,
    build_table_of_contents,
    build_sections,
    build_introduction,
    build_conclusion,
)
from content_processor import structure_content


# =============================================================================
# ТЕСТОВЫЕ ДАННЫЕ
# =============================================================================

SAMPLE_CONTENT = """# 1. Электроснабжение

## 1.1. Расчёт нагрузок

Расчётная мощность определяется по формуле:

```latex
S_{р} = \\frac{P_{р}}{\\cos \\phi}
```

Результаты расчёта сведены в таблицу:

| Показатель | Значение | Ед. изм. |
| --- | --- | --- |
| Мощность | 1500.5 | кВт |
| Ток | 32.09 | А |
| cos φ | 0.85 | — |

## 1.2. Выбор трансформатора

На основании расчётных данных выбираем трансформатор ТМ-2500/10.

- Номинальная мощность: 2500 кВА
- Напряжение ВН: 10 кВ
- Напряжение НН: 0.4 кВ
"""

SAMPLE_METADATA = DocumentMetadata(
    title="Электроснабжение жилого микрорайона",
    doc_type_label="РАСЧЁТНАЯ ЗАПИСКА",
    university="Казанский государственный энергетический университет",
    faculty="Институт электроэнергетики и электроники",
    department="Кафедра электроснабжения",
    discipline="Электроснабжение",
    author="Иванов И.И.",
    author_title="студент группы ЭС-1-21",
    city="Казань",
    year="2026"
)


# =============================================================================
# ТЕСТЫ: Титульный лист
# =============================================================================

class TestTitlePage:
    """Создание титульного листа."""

    def test_build_title_page(self):
        doc = Document()
        build_title_page(doc, SAMPLE_METADATA)
        # Проверяем наличие текста
        texts = [p.text for p in doc.paragraphs]
        full_text = " ".join(texts)
        assert "КАЗАНСКИЙ" in full_text.upper()
        assert "Иванов И.И." in full_text

    def test_title_page_with_minimal_metadata(self):
        doc = Document()
        meta = DocumentMetadata(title="Тестовая работа")
        build_title_page(doc, meta)
        texts = [p.text for p in doc.paragraphs]
        full_text = " ".join(texts)
        assert "Тестовая работа" in full_text


# =============================================================================
# ТЕСТЫ: Оглавление
# =============================================================================

class TestTableOfContents:
    """Вставка оглавления."""

    def test_build_toc(self):
        doc = Document()
        build_table_of_contents(doc)
        texts = [p.text for p in doc.paragraphs]
        full_text = " ".join(texts)
        assert "СОДЕРЖАНИЕ" in full_text


# =============================================================================
# ТЕСТЫ: Основная часть
# =============================================================================

class TestBuildSections:
    """Построение основной части."""

    def test_build_from_blocks(self):
        doc = Document()
        blocks = structure_content(SAMPLE_CONTENT)
        build_sections(doc, blocks)

        # Проверяем что заголовки есть
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 2

        # Проверяем что таблица создана
        assert len(doc.tables) >= 1


# =============================================================================
# ТЕСТЫ: Введение и Заключение
# =============================================================================

class TestIntroConclusion:
    """Введение и заключение."""

    def test_introduction(self):
        doc = Document()
        build_introduction(doc, "Целью данной работы является...")
        texts = [p.text for p in doc.paragraphs]
        full_text = " ".join(texts)
        assert "ВВЕДЕНИЕ" in full_text
        assert "Целью" in full_text

    def test_conclusion(self):
        doc = Document()
        build_conclusion(doc, "В результате работы были получены...")
        texts = [p.text for p in doc.paragraphs]
        full_text = " ".join(texts)
        assert "ЗАКЛЮЧЕНИЕ" in full_text


# =============================================================================
# ТЕСТЫ: Полный пайплайн
# =============================================================================

class TestFullPipeline:
    """Интеграционный тест: полный цикл создания документа."""

    def test_create_document_full(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            result = create_document(
                doc_type=DocType.CALCULATION_NOTE,
                raw_content=SAMPLE_CONTENT,
                metadata=SAMPLE_METADATA,
                output_path=output_path,
                sources=["ГОСТ 7.32-2017", "ГОСТ Р 7.0.100-2018"],
                introduction="Целью данной работы является расчёт электроснабжения.",
                conclusion="В результате работы определены основные параметры."
            )

            assert os.path.exists(result)
            assert os.path.getsize(result) > 0

            # Открываем и проверяем
            doc = Document(result)
            assert len(doc.paragraphs) > 10
            assert len(doc.tables) >= 1
            assert doc.styles["Normal"].font.name == "Times New Roman"

        finally:
            os.unlink(output_path)

    def test_create_minimal_document(self):
        """Минимальный документ — только текст, без источников и введения."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            result = create_document(
                doc_type=DocType.LECTURE,
                raw_content="# Лекция 1\nЗдравствуйте.",
                metadata=DocumentMetadata(title="Тест"),
                output_path=output_path
            )

            assert os.path.exists(result)
            doc = Document(result)
            assert len(doc.paragraphs) > 0

        finally:
            os.unlink(output_path)
