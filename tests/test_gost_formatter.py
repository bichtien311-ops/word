"""
Тесты для gost_formatter.py

Проверяет: создание документов, параметры страницы, стили текста,
таблицы, формулы, нумерация страниц.
"""

import os
import sys
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "execution"))

import pytest
from docx import Document
from docx.shared import Mm, Pt, Cm

from gost_formatter import (
    GostConfig,
    apply_page_setup,
    apply_base_styles,
    apply_heading_styles,
    apply_full_gost,
    create_empty_gost_document,
    create_gost_table,
    add_formula,
    add_page_numbering,
    reformat_document,
)


# =============================================================================
# ТЕСТЫ: Параметры страницы
# =============================================================================

class TestPageSetup:
    """Проверка параметров страницы по ГОСТ."""

    def test_page_margins(self):
        doc = Document()
        apply_page_setup(doc)
        section = doc.sections[0]
        # Допуск 500 EMU (~0.01мм) из-за округления в шаблоне python-docx
        assert section.left_margin == pytest.approx(Mm(30), abs=500)
        assert section.right_margin == pytest.approx(Mm(15), abs=500)
        assert section.top_margin == pytest.approx(Mm(20), abs=500)
        assert section.bottom_margin == pytest.approx(Mm(20), abs=500)

    def test_page_size_a4(self):
        doc = Document()
        apply_page_setup(doc)
        section = doc.sections[0]
        assert section.page_width == pytest.approx(Mm(210), abs=500)
        assert section.page_height == pytest.approx(Mm(297), abs=500)


# =============================================================================
# ТЕСТЫ: Базовые стили
# =============================================================================

class TestBaseStyles:
    """Проверка стиля Normal по ГОСТ."""

    def test_normal_font(self):
        doc = Document()
        apply_base_styles(doc)
        style = doc.styles["Normal"]

        assert style.font.name == "Times New Roman"
        assert style.font.size == Pt(14)
        assert style.font.bold is False

    def test_normal_paragraph_format(self):
        doc = Document()
        apply_base_styles(doc)
        style = doc.styles["Normal"]

        assert style.paragraph_format.line_spacing == 1.5
        assert style.paragraph_format.first_line_indent == pytest.approx(Cm(1.25), abs=500)


# =============================================================================
# ТЕСТЫ: Стили заголовков
# =============================================================================

class TestHeadingStyles:
    """Проверка стилей заголовков по ГОСТ."""

    def test_heading1_style(self):
        doc = Document()
        apply_heading_styles(doc)
        style = doc.styles["Heading 1"]

        assert style.font.name == "Times New Roman"
        assert style.font.size == Pt(16)
        assert style.font.bold is True
        assert style.font.all_caps is True

    def test_heading2_style(self):
        doc = Document()
        apply_heading_styles(doc)
        style = doc.styles["Heading 2"]

        assert style.font.size == Pt(14)
        assert style.font.bold is True


# =============================================================================
# ТЕСТЫ: Создание документа
# =============================================================================

class TestCreateDocument:
    """Создание ГОСТ-документа и сохранение."""

    def test_create_empty_gost_document(self):
        doc = create_empty_gost_document()
        assert doc is not None
        # Проверяем что стили применены
        assert doc.styles["Normal"].font.name == "Times New Roman"

    def test_save_document(self):
        doc = create_empty_gost_document()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            doc.save(path)
            assert os.path.exists(path)
            assert os.path.getsize(path) > 0
        finally:
            os.unlink(path)


# =============================================================================
# ТЕСТЫ: Таблицы
# =============================================================================

class TestGostTable:
    """Создание таблиц по ГОСТ."""

    def test_create_table(self):
        doc = Document()
        headers = ["Показатель", "Значение", "Ед. изм."]
        rows = [
            ["Мощность", "1500", "кВт"],
            ["Ток", "32", "А"],
        ]
        create_gost_table(doc, headers, rows, caption="Параметры", number="1.1")

        # Проверяем что таблица создана
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 3  # 1 header + 2 data rows
        assert len(table.columns) == 3

    def test_table_with_caption(self):
        doc = Document()
        create_gost_table(doc, ["A", "B"], [["1", "2"]],
                          caption="Тестовая таблица", number="2.1")
        # Подпись таблицы — это абзац перед таблицей
        found = False
        for para in doc.paragraphs:
            if "Таблица 2.1" in para.text:
                found = True
                break
        assert found


# =============================================================================
# ТЕСТЫ: Формулы
# =============================================================================

class TestFormula:
    """Добавление формул по ГОСТ."""

    def test_add_formula(self):
        doc = Document()
        add_formula(doc, "I = P / U", number="1.1")
        # Формула добавлена как абзац
        found = any("I = P / U" in p.text for p in doc.paragraphs)
        assert found

    def test_formula_with_variables(self):
        doc = Document()
        add_formula(doc, "I = P / U", number="1.1",
                    variables={"I": "ток, А", "P": "мощность, Вт"})
        texts = [p.text for p in doc.paragraphs]
        full_text = " ".join(texts)
        assert "где" in full_text


# =============================================================================
# ТЕСТЫ: Полное форматирование
# =============================================================================

class TestFullGost:
    """Комплексное ГОСТ-форматирование."""

    def test_apply_full_gost(self):
        doc = Document()
        doc.add_paragraph("Тестовый текст")
        apply_full_gost(doc)

        # Проверяем стили
        assert doc.styles["Normal"].font.name == "Times New Roman"
        assert doc.sections[0].left_margin == pytest.approx(Mm(30), abs=500)

    def test_reformat_document(self):
        # Создаём исходный документ
        doc = Document()
        doc.add_paragraph("Тестовый абзац")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = f.name
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            doc.save(input_path)
            reformat_document(input_path, output_path)

            # Открываем отформатированный
            result_doc = Document(output_path)
            assert result_doc.styles["Normal"].font.name == "Times New Roman"
            assert os.path.getsize(output_path) > 0
        finally:
            os.unlink(input_path)
            os.unlink(output_path)
